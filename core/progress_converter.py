"""
带进度回调的 PDF→Word 转换引擎（基于 pdf2docx）。
"""

import io
import logging
import os

try:
    from pdf2docx import Converter
    from pdf2docx.converter import ConversionException, MakedocxException
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_BREAK
    import fitz
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    Converter = None
    ConversionException = Exception
    MakedocxException = Exception


class ProgressConverter(Converter if Converter else object):
    """带进度回调的PDF转Word转换器"""

    def __init__(self, pdf_file=None, password=None, stream=None,
                 progress_callback=None, formula_mode=0):
        if not PDF2DOCX_AVAILABLE:
            raise RuntimeError("pdf2docx 未安装")
        super().__init__(pdf_file=pdf_file, password=password, stream=stream)
        self.progress_callback = progress_callback
        self.skipped_pages = set()
        self.formula_mode = formula_mode  # 0=普通 1=智能检测(混合) 2=全部转图片
        self.math_pages = set()
        self.image_pages = set()

    def _notify(self, phase, current, total, page_id):
        if self.progress_callback:
            self.progress_callback(phase, current, total, page_id)

    def parse_pages(self, **kwargs):
        """解析页面并回调进度"""
        logging.info(self._color_output('[3/4] Parsing pages...'))
        pages = [page for page in self._pages if not page.skip_parsing]
        total_pages = len(self._pages)
        num_pages = len(pages)
        for i, page in enumerate(pages, start=1):
            pid = page.id + 1
            self._notify('start-parse', i, num_pages, pid)
            logging.info('(%d/%d) Page %d', i, num_pages, pid)
            try:
                page.parse(**kwargs)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    logging.error('Ignore page %d due to parsing page error: %s', pid, e)
                    self.skipped_pages.add(pid)
                    self._notify('skip-parse', i, num_pages, pid)
                else:
                    raise ConversionException(f'Error when parsing page {pid}: {e}')
            finally:
                self._notify('parse', i, num_pages, pid)
        return self

    def make_docx(self, filename_or_stream=None, **kwargs):
        """生成docx并回调进度"""
        logging.info(self._color_output('[4/4] Creating pages...'))
        parsed_pages = list(filter(lambda page: page.finalized, self._pages))
        if not parsed_pages:
            raise ConversionException('No parsed pages. Please parse page first.')

        if not filename_or_stream:
            if self.filename_pdf:
                filename_or_stream = f'{self.filename_pdf[0:-len(".pdf")]}.docx'
                if os.path.exists(filename_or_stream):
                    os.remove(filename_or_stream)
            else:
                raise ConversionException('Please specify a docx file name or a file-like object to write.')

        docx_file = Document()
        num_pages = len(parsed_pages)
        for i, page in enumerate(parsed_pages, start=1):
            if not page.finalized:
                continue
            pid = page.id + 1
            self._notify('start-make', i, num_pages, pid)
            logging.info('(%d/%d) Page %d', i, num_pages, pid)
            try:
                if self.formula_mode == 2 and page.id in self.math_pages:
                    self._render_page_as_image(docx_file, page.id, i > 1)
                    self.image_pages.add(pid)
                    logging.info('Page %d rendered as whole-page image', pid)
                else:
                    page.make_docx(docx_file)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    logging.error('Ignore page %d due to making page error: %s', pid, e)
                    self.skipped_pages.add(pid)
                    self._notify('skip-make', i, num_pages, pid)
                else:
                    raise MakedocxException(f'Error when make page {pid}: {e}')
            finally:
                self._notify('make', i, num_pages, pid)

        docx_file.save(filename_or_stream)

    def _render_page_as_image(self, docx_file, page_id, add_page_break=True):
        """将PDF页面渲染为高清图片并插入Word文档（整页模式）"""
        pdf_page = self.fitz_doc[page_id]
        dpi = 300
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = pdf_page.get_pixmap(matrix=mat)
        img_stream = io.BytesIO(pix.tobytes("png"))
        page_width_inches = pdf_page.rect.width / 72.0
        target_width = min(page_width_inches, 6.3)
        if add_page_break and len(docx_file.paragraphs) > 0:
            run = docx_file.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
        docx_file.add_picture(img_stream, width=Inches(target_width))
