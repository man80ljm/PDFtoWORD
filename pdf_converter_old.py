"""
PDF转Word转换工具
使用tkinter构建的图形界面应用程序
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
from datetime import datetime
import threading
import logging
import time
import json
import shutil
import sys
import io
import base64
import re

# 配置日志输出到文件（方便调试API问题）
_log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)) if not getattr(sys, 'frozen', False)
                         else os.path.dirname(sys.executable), 'pdf_converter.log')
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler(_log_file, encoding='utf-8', mode='a'),
        logging.StreamHandler()
    ]
)

# API调用和公式转换相关库
try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

try:
    import latex2mathml.converter
    from lxml import etree
    LATEX2OMML_AVAILABLE = True
except ImportError:
    LATEX2OMML_AVAILABLE = False

# PDF转换相关库
try:
    from pdf2docx import Converter
    from pdf2docx.converter import ConversionException, MakedocxException
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    import fitz
    PDF2DOCX_AVAILABLE = True

    # 已知的数学字体名称模式（用于公式页面检测）
    MATH_FONT_PATTERNS = [
        # Computer Modern (LaTeX经典)
        'CMMI', 'CMSY', 'CMEX', 'CMR', 'CMTI', 'CMBX', 'CMSS',
        # AMS 字体
        'MSAM', 'MSBM', 'EUSM', 'EUFM', 'EURM', 'EUBM',
        # Latin Modern (现代LaTeX默认)
        'LatinModernMath', 'LMMath', 'LatinModern-Math', 'LMRoman', 'LMSans',
        # STIX / XITS
        'STIX', 'XITS', 'STIXMath', 'XITSMath',
        # Cambria Math (Office常用)
        'CambriaMath', 'Cambria Math', 'Cambria-Math',
        # Libertinus / Linux Libertine
        'LibertinusMath', 'Libertinus Math', 'LinuxLibertine',
        # TeX Gyre
        'TeXGyre', 'TeX Gyre', 'TexGyreMath',
        # Fira Math
        'FiraMath', 'Fira Math',
        # Asana Math
        'AsanaMath', 'Asana-Math', 'Asana Math',
        # DejaVu Math
        'DejaVuMath', 'DejaVu Math',
        # Garamond Math
        'GaramondMath', 'Garamond-Math',
        # Symbol / MathType
        'Symbol', 'MT Extra', 'MT Symbol',
        'Mathematica', 'MathematicalPi',
        'Euclid',
        # 其他
        'RSFS', 'WASY', 'LASY',
    ]

    # 字体名称中包含这些关键词的也视为数学字体
    MATH_FONT_KEYWORDS = ['math', 'symbol', 'cmmi', 'cmsy', 'cmex']

    def detect_math_pages(fitz_doc, start=0, end=None):
        """检测包含数学公式的页面（通过分析字体、CID字体、Type3字体和特殊字符）"""
        if end is None:
            end = len(fitz_doc)
        math_pages = set()
        for page_idx in range(start, end):
            page = fitz_doc[page_idx]
            fonts = page.get_fonts()
            has_math_font = False
            for font in fonts:
                # font: (xref, ext, type, basefont, name, encoding)
                font_type = font[2] if len(font) > 2 else ""
                font_basefont = font[3] if len(font) > 3 else ""
                # 去掉子集前缀 (如 "ABCDEF+")
                clean_name = font_basefont
                if '+' in clean_name:
                    clean_name = clean_name.split('+', 1)[1]
                clean_lower = clean_name.lower().replace('-', '').replace(' ', '')
                # 1. 精确匹配已知数学字体模式
                for pat in MATH_FONT_PATTERNS:
                    if pat.lower().replace('-', '').replace(' ', '') in clean_lower:
                        has_math_font = True
                        break
                if has_math_font:
                    break
                # 2. 关键词匹配
                for kw in MATH_FONT_KEYWORDS:
                    if kw in clean_lower:
                        has_math_font = True
                        break
                if has_math_font:
                    break
                # 3. Type3 字体常用于嵌入的数学符号
                if font_type == 'Type3':
                    has_math_font = True
                    break
            if has_math_font:
                math_pages.add(page_idx)
        return math_pages

    def _is_math_font(font_name):
        """判断字体名是否为数学字体"""
        if not font_name:
            return False
        clean = font_name
        if '+' in clean:
            clean = clean.split('+', 1)[1]
        clean_lower = clean.lower().replace('-', '').replace(' ', '')
        for pat in MATH_FONT_PATTERNS:
            if pat.lower().replace('-', '').replace(' ', '') in clean_lower:
                return True
        for kw in MATH_FONT_KEYWORDS:
            if kw in clean_lower:
                return True
        return False

    def _has_math_unicode(text):
        """检查文本是否包含需要规范化的数学Unicode字符"""
        for c in text:
            cp = ord(c)
            if 0x1D400 <= cp <= 0x1D7FF:  # Mathematical Alphanumeric Symbols
                return True
            if cp == 0x210E:  # PLANCK CONSTANT
                return True
        return False

    def _is_display_equation(block):
        """判断一个块是否为独立的行间公式（大部分为数学字体，不含CJK字符）"""
        if block.get("type") != 0:
            return False
        total_chars = 0
        math_chars = 0
        cjk_chars = 0
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                text = span.get("text", "").strip()
                font = span.get("font", "")
                is_math = _is_math_font(font)
                for c in text:
                    if c.isspace():
                        continue
                    total_chars += 1
                    if is_math:
                        math_chars += 1
                    if 0x4E00 <= ord(c) <= 0x9FFF:
                        cjk_chars += 1
        if total_chars < 2:
            return False
        # 独立公式：大部分内容为数学字体，且不含中文字符
        return math_chars / total_chars > 0.5 and cjk_chars == 0

    def _get_block_text(block):
        """提取块中所有span的文本"""
        parts = []
        for line in block.get("lines", []):
            line_parts = []
            for span in line.get("spans", []):
                line_parts.append(span.get("text", ""))
            parts.append("".join(line_parts))
        return " ".join(parts).strip()

    def _normalize_math_unicode(text):
        """将 Unicode 数学字母数字符号转为普通字符，使 Word 能正确显示。
        例如: 𝑓(U+1D453) → f, 𝑥(U+1D465) → x, 𝜋(U+1D70B) → π"""
        if not text:
            return text
        result = []
        for c in text:
            cp = ord(c)
            mapped = _MAP_MATH_CHAR(cp)
            result.append(mapped)
        return ''.join(result)

    def _MAP_MATH_CHAR(cp):
        """将数学Unicode码点映射为普通可显示字符"""
        # Mathematical Italic Small (U+1D44E - U+1D467) → a-z
        if 0x1D44E <= cp <= 0x1D467:
            return chr(ord('a') + cp - 0x1D44E)
        # Mathematical Italic Capital (U+1D434 - U+1D44D) → A-Z
        if 0x1D434 <= cp <= 0x1D44D:
            return chr(ord('A') + cp - 0x1D434)
        # Mathematical Bold Small (U+1D41A - U+1D433) → a-z
        if 0x1D41A <= cp <= 0x1D433:
            return chr(ord('a') + cp - 0x1D41A)
        # Mathematical Bold Capital (U+1D400 - U+1D419) → A-Z
        if 0x1D400 <= cp <= 0x1D419:
            return chr(ord('A') + cp - 0x1D400)
        # Mathematical Bold Italic Small (U+1D482 - U+1D49B) → a-z
        if 0x1D482 <= cp <= 0x1D49B:
            return chr(ord('a') + cp - 0x1D482)
        # Mathematical Bold Italic Capital (U+1D468 - U+1D481) → A-Z
        if 0x1D468 <= cp <= 0x1D481:
            return chr(ord('A') + cp - 0x1D468)
        # Mathematical Sans-Serif variants
        if 0x1D5A0 <= cp <= 0x1D5B9:  # sans capital
            return chr(ord('A') + cp - 0x1D5A0)
        if 0x1D5BA <= cp <= 0x1D5D3:  # sans small
            return chr(ord('a') + cp - 0x1D5BA)
        # Mathematical Italic Greek Small (U+1D6FC - U+1D714) → α-ω
        _GREEK_LOWER = 'αβγδεζηθικλμνξοπρςστυφχψω'
        if 0x1D6FC <= cp <= 0x1D714:
            idx = cp - 0x1D6FC
            if idx < len(_GREEK_LOWER):
                return _GREEK_LOWER[idx]
        # Mathematical Italic Greek Capital (U+1D6E2 - U+1D6FA) → Α-Ω
        _GREEK_UPPER = 'ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡ΢ΣΤΥΦΧΨΩ'
        if 0x1D6E2 <= cp <= 0x1D6FA:
            idx = cp - 0x1D6E2
            if idx < len(_GREEK_UPPER):
                return _GREEK_UPPER[idx]
        # Mathematical Bold Greek Small (U+1D736 - U+1D74E)
        if 0x1D736 <= cp <= 0x1D74E:
            idx = cp - 0x1D736
            if idx < len(_GREEK_LOWER):
                return _GREEK_LOWER[idx]
        # Mathematical Bold Greek Capital (U+1D71C - U+1D734)
        if 0x1D71C <= cp <= 0x1D734:
            idx = cp - 0x1D71C
            if idx < len(_GREEK_UPPER):
                return _GREEK_UPPER[idx]
        # 数学运算符映射
        _MATH_OPERATORS = {
            0x2212: '-',   # MINUS SIGN → -
            0x2032: "'",   # PRIME → '
            0x2033: "''",  # DOUBLE PRIME
            0x2190: '←', 0x2192: '→', 0x21D2: '⇒', 0x21D0: '⇐',
            0x2260: '≠', 0x2264: '≤', 0x2265: '≥',
            0x222B: '∫', 0x2211: '∑', 0x220F: '∏',
            0x221A: '√', 0x221E: '∞', 0x2202: '∂',
            0x210E: 'h',  # PLANCK CONSTANT → h
        }
        if cp in _MATH_OPERATORS:
            return _MATH_OPERATORS[cp]
        return chr(cp)

    class ProgressConverter(Converter):
        """带进度回调的PDF转Word转换器"""

        def __init__(self, pdf_file: str = None, password: str = None, stream: bytes = None, progress_callback=None, formula_mode=0):
            super().__init__(pdf_file=pdf_file, password=password, stream=stream)
            self.progress_callback = progress_callback
            self.skipped_pages = set()
            self.formula_mode = formula_mode  # 0=普通 1=智能检测(混合) 2=全部转图片
            self.math_pages = set()
            self.image_pages = set()

        def _notify(self, phase: str, current: int, total: int, page_id: int):
            if self.progress_callback:
                self.progress_callback(phase, current, total, page_id)

        def parse_pages(self, **kwargs):
            """解析页面并回调进度"""
            logging.info(self._color_output('[3/4] Parsing pages...'))

            pages = [page for page in self._pages if not page.skip_parsing]
            total_pages = len(self._pages)
            num_pages = len(pages)
            for i, page in enumerate(pages, start=1):
                pid = page.id + 1
                self._notify('start-parse', i, num_pages, pid)
                logging.info('(%d/%d) Page %d', i, num_pages, pid)
                try:
                    page.parse(**kwargs)
                except Exception as e:
                    if not kwargs['debug'] and kwargs['ignore_page_error']:
                        logging.error('Ignore page %d due to parsing page error: %s', pid, e)
                        self.skipped_pages.add(pid)
                        self._notify('skip-parse', i, num_pages, pid)
                    else:
                        raise ConversionException(f'Error when parsing page {pid}: {e}')
                finally:
                    self._notify('parse', i, num_pages, pid)

            return self

        def make_docx(self, filename_or_stream=None, **kwargs):
            """生成docx并回调进度"""
            logging.info(self._color_output('[4/4] Creating pages...'))

            parsed_pages = list(filter(lambda page: page.finalized, self._pages))
            if not parsed_pages:
                raise ConversionException('No parsed pages. Please parse page first.')

            if not filename_or_stream:
                if self.filename_pdf:
                    filename_or_stream = f'{self.filename_pdf[0:-len(".pdf")]}.docx'
                    if os.path.exists(filename_or_stream):
                        os.remove(filename_or_stream)
                else:
                    raise ConversionException('Please specify a docx file name or a file-like object to write.')

            docx_file = Document()
            num_pages = len(parsed_pages)
            for i, page in enumerate(parsed_pages, start=1):
                if not page.finalized:
                    continue
                pid = page.id + 1
                self._notify('start-make', i, num_pages, pid)
                logging.info('(%d/%d) Page %d', i, num_pages, pid)
                try:
                    if self.formula_mode == 2 and page.id in self.math_pages:
                        # 模式2：全部转图片
                        self._render_page_as_image(docx_file, page.id, i > 1)
                        self.image_pages.add(pid)
                        logging.info('Page %d rendered as whole-page image', pid)
                    else:
                        # 模式0和模式1：pdf2docx正常转换（模式1会在后处理中修复）
                        page.make_docx(docx_file)
                except Exception as e:
                    if not kwargs['debug'] and kwargs['ignore_page_error']:
                        logging.error('Ignore page %d due to making page error: %s', pid, e)
                        self.skipped_pages.add(pid)
                        self._notify('skip-make', i, num_pages, pid)
                    else:
                        raise MakedocxException(f'Error when make page {pid}: {e}')
                finally:
                    self._notify('make', i, num_pages, pid)

            docx_file.save(filename_or_stream)

        def _render_page_as_image(self, docx_file, page_id, add_page_break=True):
            """将PDF页面渲染为高清图片并插入Word文档（整页模式）"""
            pdf_page = self.fitz_doc[page_id]
            dpi = 300
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = pdf_page.get_pixmap(matrix=mat)
            img_stream = io.BytesIO(pix.tobytes("png"))
            page_width_inches = pdf_page.rect.width / 72.0
            target_width = min(page_width_inches, 6.3)
            if add_page_break and len(docx_file.paragraphs) > 0:
                run = docx_file.add_paragraph().add_run()
                run.add_break(WD_BREAK.PAGE)
            docx_file.add_picture(img_stream, width=Inches(target_width))

except ImportError:
    PDF2DOCX_AVAILABLE = False
    ProgressConverter = None

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ============================================================
# API 辅助类
# ============================================================

def _simple_encrypt(text):
    """简单混淆存储（非安全加密，仅避免明文）"""
    if not text:
        return ""
    return base64.b64encode(text.encode('utf-8')).decode('utf-8')

def _simple_decrypt(encoded):
    """解码简单混淆"""
    if not encoded:
        return ""
    try:
        return base64.b64decode(encoded.encode('utf-8')).decode('utf-8')
    except Exception:
        return encoded  # 兼容旧版明文


class BaiduOCRClient:
    """百度OCR API客户端 - 支持通用文字识别和公式识别"""

    TOKEN_URL = "https://aip.baidubce.com/oauth/2.0/token"
    OCR_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"
    FORMULA_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/formula"

    def __init__(self, api_key, secret_key):
        self.api_key = api_key
        self.secret_key = secret_key
        self._access_token = None
        self._token_time = 0

    def _get_access_token(self):
        """获取百度API access_token（有效期30天，自动缓存）"""
        if self._access_token and (time.time() - self._token_time) < 86400 * 25:
            return self._access_token
        params = {
            "grant_type": "client_credentials",
            "client_id": self.api_key,
            "client_secret": self.secret_key,
        }
        resp = requests.post(self.TOKEN_URL, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        if "access_token" not in data:
            raise RuntimeError(f"百度API认证失败: {data.get('error_description', data)}")
        self._access_token = data["access_token"]
        self._token_time = time.time()
        return self._access_token

    def test_connection(self):
        """测试API连接是否可用"""
        try:
            self._get_access_token()
            return True, "连接成功"
        except Exception as e:
            return False, str(e)

    @staticmethod
    def _compress_image(image_bytes, max_size_bytes=3 * 1024 * 1024):
        """将图片压缩为JPEG格式，确保不超过百度API的大小限制"""
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        # 先尝试高质量JPEG
        for quality in [90, 80, 65, 50]:
            buf = io.BytesIO()
            img.save(buf, 'JPEG', quality=quality)
            jpg_bytes = buf.getvalue()
            b64_len = len(base64.b64encode(jpg_bytes))
            if b64_len <= max_size_bytes:
                logging.info(f'Image compressed: {len(image_bytes)//1024}KB→{len(jpg_bytes)//1024}KB (q={quality}, b64={b64_len//1024}KB)')
                return jpg_bytes
        # 如果还是太大，缩小尺寸
        w, h = img.size
        img = img.resize((w // 2, h // 2), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, 'JPEG', quality=70)
        return buf.getvalue()

    def recognize_text(self, image_bytes):
        """通用文字识别（高精度版），返回文字行列表"""
        token = self._get_access_token()
        compressed = self._compress_image(image_bytes)
        img_b64 = base64.b64encode(compressed).decode()
        logging.info(f'OCR text request: image base64 size = {len(img_b64)//1024} KB')
        headers = {"Content-Type": "application/x-www-form-urlencoded"}
        data = {
            "image": img_b64,
            "language_type": "CHN_ENG",
            "detect_direction": "true",
            "paragraph": "true",
        }
        resp = requests.post(
            f"{self.OCR_URL}?access_token={token}",
            headers=headers, data=data, timeout=60
        )
        resp.raise_for_status()
        result = resp.json()
        logging.info(f'OCR text response keys: {list(result.keys())}, words_num: {result.get("words_result_num", 0)}')
        if "error_code" in result:
            raise RuntimeError(f"OCR识别失败[{result.get('error_code')}]: {result.get('error_msg', result)}")
        words = []
        for item in result.get("words_result", []):
            words.append(item.get("words", ""))
        return words

    def recognize_formula(self, image_bytes):
        """公式识别，返回 LaTeX 字符串列表"""
        token = self._get_access_token()
        compressed = self._compress_image(image_bytes)
        img_b64 = base64.b64encode(compressed).decode()
        logging.info(f'Formula request: image base64 size = {len(img_b64)//1024} KB')
        headers = {"Content-Type": "application/x-www-form-urlencoded"}
        data = {
            "image": img_b64,
            "recognize_granularity": "big",
        }
        resp = requests.post(
            f"{self.FORMULA_URL}?access_token={token}",
            headers=headers, data=data, timeout=60
        )
        resp.raise_for_status()
        result = resp.json()
        logging.info(f'Formula response keys: {list(result.keys())}')
        if "error_code" in result:
            raise RuntimeError(f"公式识别失败[{result.get('error_code')}]: {result.get('error_msg', result)}")
        formulas = []
        # 百度API可能返回 words_result 或 formulas_result，两个都尝试
        formula_items = result.get("formulas_result", result.get("words_result", []))
        for item in formula_items:
            text = item.get("words", "")
            if text:
                formulas.append(text)
                logging.info(f'  Formula detected: {text[:80]}')
        if not formulas:
            logging.info(f'  No formulas found in response: {str(result)[:200]}')
        return formulas


def latex_to_omml(latex_str, xslt_path=None):
    """将LaTeX公式转为Word OMML XML元素。
    需要 latex2mathml 和 lxml，以及 MML2OMML.XSL（Office自带或内嵌）。
    """
    if not LATEX2OMML_AVAILABLE:
        return None

    # 清理LaTeX（去掉可能的$包裹）
    latex_clean = latex_str.strip()
    for prefix in ['$$', '$', '\\[', '\\(']:
        if latex_clean.startswith(prefix):
            latex_clean = latex_clean[len(prefix):]
    for suffix in ['$$', '$', '\\]', '\\)']:
        if latex_clean.endswith(suffix):
            latex_clean = latex_clean[:-len(suffix)]
    latex_clean = latex_clean.strip()
    if not latex_clean:
        return None

    # LaTeX → MathML
    try:
        mathml_str = latex2mathml.converter.convert(latex_clean)
    except Exception as e:
        logging.warning(f"LaTeX→MathML转换失败: {e}, 原始: {latex_clean}")
        return None

    # MathML → OMML via XSLT
    # 尝试查找 MML2OMML.XSL
    if xslt_path is None:
        candidate_paths = [
            r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
            r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\root\Office15\MML2OMML.XSL",
        ]
        for p in candidate_paths:
            if os.path.exists(p):
                xslt_path = p
                break

    if xslt_path is None or not os.path.exists(xslt_path):
        logging.warning("未找到 MML2OMML.XSL，无法将MathML转为OMML")
        return None

    try:
        with open(xslt_path, 'rb') as f:
            xslt_doc = etree.parse(f)
        transform = etree.XSLT(xslt_doc)
        mathml_doc = etree.fromstring(mathml_str.encode())
        omml_result = transform(mathml_doc)
        omml_element = omml_result.getroot()
        return omml_element
    except Exception as e:
        logging.warning(f"MathML→OMML转换失败: {e}")
        return None


def insert_omml_to_paragraph(paragraph, omml_element):
    """将OMML公式元素插入到Word段落中"""
    from lxml import etree
    MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    # 将 lxml element 转为 python-docx 兼容的 element
    omml_str = etree.tostring(omml_element)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import copy
    # 解析为 python-docx 兼容的 XML
    omml_parsed = etree.fromstring(omml_str)
    # 直接 append 到 paragraph 的 XML
    paragraph._element.append(omml_parsed)

class PDFConverterApp:
    """PDF转换工具主应用类"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("程新伟专属转换器 - PDF转Word")
        self.root.geometry("500x580")
        self.root.resizable(False, False)
        
        # 设置应用图标（如果可用）
        try:
            self.root.iconbitmap('icon.ico')
        except:
            pass
        
        # 变量
        self.selected_file = tk.StringVar()
        self.status_message = tk.StringVar(value="就绪")
        self.total_pages = 0
        self.total_steps = 0
        self.start_time = None
        self.current_page_id = None
        self.current_page_index = None
        self.current_page_total = None
        self.current_phase = None
        self.page_start_time = None
        self.page_timeout_seconds = 60
        self.page_timer_job = None
        self.current_eta_text = ""
        self.base_status_text = ""
        self.conversion_active = False
        self.page_start_var = tk.StringVar()
        self.page_end_var = tk.StringVar()
        self.formula_mode_var = tk.IntVar(value=0)  # 保留兼容
        self.title_text_var = tk.StringVar(value="程新伟专属转换器")
        self.settings_path = os.path.join(self.get_app_dir(), "settings.json")
        self.bg_image_path = None
        self.bg_image = None
        self.bg_pil = None
        self.bg_label = None
        self.panel_opacity_var = tk.DoubleVar(value=85.0)
        self.panel_padding = 20
        self.panel_image = None
        self.panel_canvas = None
        self.panel_image_id = None
        self.resize_job = None
        self.panel_resize_job = None

        # 功能选择
        self.current_function_var = tk.StringVar(value="PDF转Word")
        self.selected_files_list = []  # 批量文件列表

        # PDF转图片选项
        self.image_dpi_var = tk.StringVar(value="200")
        self.image_format_var = tk.StringVar(value="PNG")

        # 新增：OCR & 公式识别选项
        self.ocr_enabled_var = tk.BooleanVar(value=False)
        self.formula_api_enabled_var = tk.BooleanVar(value=False)

        # API 配置（从设置加载）
        self.api_provider = "baidu"  # baidu
        self.baidu_api_key = ""
        self.baidu_secret_key = ""
        self.xslt_path = None  # MML2OMML.XSL 路径
        self._baidu_client = None  # 缓存的 BaiduOCRClient
        
        # 创建UI
        self.create_ui()

        # 加载设置
        self.load_settings()
        
        # 检查依赖
        self.check_dependencies()
    
    def create_ui(self):
        """创建用户界面 - Canvas直绘实现透明面板"""

        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        self.panel_canvas = tk.Canvas(self.root, highlightthickness=0, bd=0)
        self.panel_canvas.grid(
            row=0, column=0, sticky="nsew",
            padx=self.panel_padding, pady=self.panel_padding
        )

        # 设置按钮
        self.settings_btn = tk.Button(
            self.panel_canvas, text="⚙", font=("Microsoft YaHei", 12),
            relief=tk.FLAT, padx=4, cursor='hand2',
            command=self.open_settings_window
        )
        self.cv_settings = self.panel_canvas.create_window(5, 5, window=self.settings_btn, anchor="nw")

        # 标题（透明背景）
        self.cv_title = self.panel_canvas.create_text(
            0, 35, text=self.title_text_var.get(),
            font=("Microsoft YaHei", 26, "bold"), anchor="n"
        )
        self.title_text_var.trace_add("write", self._on_title_var_changed)

        # 副标题区 → 功能选择器
        func_frame = tk.Frame(self.panel_canvas)
        tk.Label(func_frame, text="功能:", font=("Microsoft YaHei", 10, "bold")).pack(side=tk.LEFT)
        self.func_combo = ttk.Combobox(
            func_frame, textvariable=self.current_function_var,
            values=["PDF转Word", "PDF转图片"],
            state='readonly', font=("Microsoft YaHei", 10), width=14
        )
        self.func_combo.pack(side=tk.LEFT, padx=(8, 0))
        self.func_combo.bind("<<ComboboxSelected>>", self._on_function_changed)
        self.cv_subtitle = self.panel_canvas.create_window(
            0, 75, window=func_frame, anchor="n"
        )

        # 分区标题：选择PDF文件（透明背景）
        self.cv_section1 = self.panel_canvas.create_text(
            15, 105, text="选择PDF文件",
            font=("Microsoft YaHei", 11, "bold"), anchor="nw"
        )

        # 文件输入框 + 浏览按钮
        file_frame = tk.Frame(self.panel_canvas)
        self.file_entry = tk.Entry(
            file_frame, textvariable=self.selected_file,
            font=("Microsoft YaHei", 10), state='readonly'
        )
        self.file_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
        tk.Button(
            file_frame, text="浏览...", command=self.browse_file,
            font=("Microsoft YaHei", 10), padx=20, cursor='hand2'
        ).pack(side=tk.LEFT, padx=(10, 0), ipady=6)
        self.cv_file_frame = self.panel_canvas.create_window(
            15, 130, window=file_frame, anchor="nw", width=1
        )

        # 分区标题：页范围（透明背景）
        self.cv_section2 = self.panel_canvas.create_text(
            15, 185, text="页范围（可选）",
            font=("Microsoft YaHei", 11, "bold"), anchor="nw"
        )

        # 页范围输入
        range_frame = tk.Frame(self.panel_canvas)
        tk.Label(range_frame, text="起始页:", font=("Microsoft YaHei", 10)).pack(side=tk.LEFT)
        tk.Entry(range_frame, textvariable=self.page_start_var, width=6,
                 font=("Microsoft YaHei", 10)).pack(side=tk.LEFT, padx=(6, 20))
        tk.Label(range_frame, text="结束页:", font=("Microsoft YaHei", 10)).pack(side=tk.LEFT)
        tk.Entry(range_frame, textvariable=self.page_end_var, width=6,
                 font=("Microsoft YaHei", 10)).pack(side=tk.LEFT, padx=(6, 20))
        tk.Label(range_frame, text="留空表示全部页（页码从1开始）",
                 font=("Microsoft YaHei", 9)).pack(side=tk.LEFT)
        self.cv_range_frame = self.panel_canvas.create_window(
            15, 210, window=range_frame, anchor="nw"
        )

        # 转换选项区（Word模式）
        self.word_options_frame = tk.Frame(self.panel_canvas)
        tk.Label(self.word_options_frame, text="转换选项:", font=("Microsoft YaHei", 9, "bold")).pack(side=tk.LEFT)
        self.ocr_cb = tk.Checkbutton(
            self.word_options_frame, text="OCR识别(扫描件)",
            variable=self.ocr_enabled_var, font=("Microsoft YaHei", 9),
            command=self._on_option_changed
        )
        self.ocr_cb.pack(side=tk.LEFT, padx=(8, 0))
        self.formula_cb = tk.Checkbutton(
            self.word_options_frame, text="公式智能识别",
            variable=self.formula_api_enabled_var, font=("Microsoft YaHei", 9),
            command=self._on_option_changed
        )
        self.formula_cb.pack(side=tk.LEFT, padx=(8, 0))
        self.cv_formula_frame = self.panel_canvas.create_window(
            15, 245, window=self.word_options_frame, anchor="nw"
        )

        # 转换选项区（图片模式）
        self.image_options_frame = tk.Frame(self.panel_canvas)
        tk.Label(self.image_options_frame, text="输出设置:", font=("Microsoft YaHei", 9, "bold")).pack(side=tk.LEFT)
        tk.Label(self.image_options_frame, text="DPI:", font=("Microsoft YaHei", 9)).pack(side=tk.LEFT, padx=(10, 0))
        dpi_combo = ttk.Combobox(
            self.image_options_frame, textvariable=self.image_dpi_var,
            values=["72", "150", "200", "300", "600"],
            width=5, font=("Microsoft YaHei", 9)
        )
        dpi_combo.pack(side=tk.LEFT, padx=(4, 0))
        tk.Label(self.image_options_frame, text="格式:", font=("Microsoft YaHei", 9)).pack(side=tk.LEFT, padx=(14, 0))
        fmt_combo = ttk.Combobox(
            self.image_options_frame, textvariable=self.image_format_var,
            values=["PNG", "JPEG"],
            state='readonly', width=6, font=("Microsoft YaHei", 9)
        )
        fmt_combo.pack(side=tk.LEFT, padx=(4, 0))
        self.cv_image_options = self.panel_canvas.create_window(
            15, 245, window=self.image_options_frame, anchor="nw"
        )
        # 默认隐藏图片选项
        self.panel_canvas.itemconfigure(self.cv_image_options, state='hidden')

        # API状态提示
        self.cv_api_hint = self.panel_canvas.create_text(
            15, 270, text="", font=("Microsoft YaHei", 8), anchor="nw", fill="#888888"
        )

        # 进度条
        self.progress_bar = ttk.Progressbar(self.panel_canvas, mode='determinate')
        self.cv_progress_bar = self.panel_canvas.create_window(
            20, 290, window=self.progress_bar, anchor="nw", width=1, height=25
        )

        # 进度文本（透明背景）
        self.cv_progress_text = self.panel_canvas.create_text(
            0, 325, text="", font=("Microsoft YaHei", 9), anchor="n"
        )

        # 转换 / 清除按钮
        btn_frame = tk.Frame(self.panel_canvas)
        self.convert_btn = tk.Button(
            btn_frame, text="开始转换", command=self.start_conversion,
            font=("Microsoft YaHei", 12, "bold"), padx=40, pady=12, cursor='hand2'
        )
        self.convert_btn.pack(side=tk.LEFT, expand=True, padx=5)
        tk.Button(
            btn_frame, text="清除", command=self.clear_selection,
            font=("Microsoft YaHei", 12), padx=40, pady=12, cursor='hand2'
        ).pack(side=tk.LEFT, expand=True, padx=5)
        self.cv_btn_frame = self.panel_canvas.create_window(
            0, 370, window=btn_frame, anchor="n"
        )

        # 状态栏文字（透明背景）
        self.cv_status_text = self.panel_canvas.create_text(
            15, 0, text=self.status_message.get(),
            font=("Microsoft YaHei", 9), anchor="sw"
        )
        self.status_message.trace_add("write", self._on_status_var_changed)

        # 绑定事件
        self.root.bind("<Configure>", self.on_root_resize)
        self.panel_canvas.bind("<Configure>", self.on_panel_resize)
        self.root.after(50, self.refresh_layout)
    
    def _on_title_var_changed(self, *args):
        """标题变量变化时更新Canvas文字"""
        if self.panel_canvas:
            self.panel_canvas.itemconfigure(self.cv_title, text=self.title_text_var.get())

    def _on_status_var_changed(self, *args):
        """状态变量变化时更新Canvas文字"""
        if self.panel_canvas:
            self.panel_canvas.itemconfigure(self.cv_status_text, text=self.status_message.get())

    def set_progress_text(self, text):
        """更新进度文本"""
        if self.panel_canvas:
            self.panel_canvas.itemconfigure(self.cv_progress_text, text=text)

    def layout_canvas(self):
        """根据Canvas尺寸重新布局所有元素"""
        w = self.panel_canvas.winfo_width()
        h = self.panel_canvas.winfo_height()
        if w <= 1 or h <= 1:
            return
        cx = w // 2
        self.panel_canvas.coords(self.cv_title, cx, 35)
        self.panel_canvas.coords(self.cv_subtitle, cx, 75)
        self.panel_canvas.coords(self.cv_section1, 15, 105)
        self.panel_canvas.coords(self.cv_file_frame, 15, 130)
        self.panel_canvas.itemconfigure(self.cv_file_frame, width=w - 30)
        self.panel_canvas.coords(self.cv_section2, 15, 185)
        self.panel_canvas.coords(self.cv_range_frame, 15, 210)
        self.panel_canvas.coords(self.cv_formula_frame, 15, 245)
        self.panel_canvas.coords(self.cv_image_options, 15, 245)
        self.panel_canvas.coords(self.cv_api_hint, 15, 270)
        self.panel_canvas.coords(self.cv_progress_bar, 20, 290)
        self.panel_canvas.itemconfigure(self.cv_progress_bar, width=w - 40)
        self.panel_canvas.coords(self.cv_progress_text, cx, 325)
        self.panel_canvas.coords(self.cv_btn_frame, cx, 370)
        self.panel_canvas.coords(self.cv_status_text, 15, h - 10)
    
    def check_dependencies(self):
        """检查依赖库"""
        missing = []
        
        if not PDF2DOCX_AVAILABLE:
            missing.append("pdf2docx")
        if missing:
            msg = f"警告：以下依赖库未安装：\n{', '.join(missing)}\n\n请运行: pip install {' '.join(missing)}"
            self.status_message.set(f"缺少依赖库: {', '.join(missing)}")
            messagebox.showwarning("缺少依赖", msg)
    
    def browse_file(self):
        """浏览并选择PDF文件（图片模式支持多选）"""
        func = self.current_function_var.get()
        if func == "PDF转图片":
            filenames = filedialog.askopenfilenames(
                title="选择PDF文件（可多选）",
                filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
            )
            if filenames:
                self.selected_files_list = list(filenames)
                count = len(self.selected_files_list)
                if count == 1:
                    self.selected_file.set(filenames[0])
                    self.status_message.set(f"已选择: {os.path.basename(filenames[0])}")
                else:
                    self.selected_file.set(f"已选择 {count} 个PDF文件")
                    names = ", ".join(os.path.basename(f) for f in filenames[:3])
                    if count > 3:
                        names += f" 等共{count}个"
                    self.status_message.set(f"已选择: {names}")
        else:
            filename = filedialog.askopenfilename(
                title="选择PDF文件",
                filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
            )
            if filename:
                self.selected_file.set(filename)
                self.selected_files_list = [filename]
                self.status_message.set(f"已选择: {os.path.basename(filename)}")
    
    def clear_selection(self):
        """清除选择"""
        self.selected_file.set("")
        self.selected_files_list = []
        self.progress_bar['value'] = 0
        self.set_progress_text("")
        self.status_message.set("就绪")
        self.total_pages = 0
        self.total_steps = 0
        self.start_time = None
        self.current_page_id = None
        self.current_page_index = None
        self.current_page_total = None
        self.current_phase = None
        self.page_start_time = None
        self.current_eta_text = ""
        self.base_status_text = ""
        self.conversion_active = False
        self.page_start_var.set("")
        self.page_end_var.set("")
    
    def start_conversion(self):
        """开始转换"""
        func = self.current_function_var.get()

        if func == "PDF转图片":
            # 图片模式：检查文件列表
            if not self.selected_files_list:
                messagebox.showwarning("提示", "请先选择PDF文件！")
                return
            for f in self.selected_files_list:
                if not os.path.exists(f):
                    messagebox.showerror("错误", f"文件不存在：\n{f}")
                    return
        else:
            # Word模式：检查单文件
            if not self.selected_file.get():
                messagebox.showwarning("提示", "请先选择一个PDF文件！")
                return
            if not os.path.exists(self.selected_file.get()):
                messagebox.showerror("错误", "选择的文件不存在！")
                return
        
        # 禁用转换按钮
        self.convert_btn.config(state=tk.DISABLED)
        self.conversion_active = True
        self.current_page_id = None
        self.current_page_index = None
        self.current_page_total = None
        self.current_phase = None
        self.page_start_time = None
        self.current_eta_text = ""
        self.base_status_text = ""
        self.start_page_timer()
        
        # 在新线程中执行转换
        thread = threading.Thread(target=self.perform_conversion)
        thread.daemon = True
        thread.start()
    
    def perform_conversion(self):
        """执行转换（在后台线程中）"""
        try:
            func = self.current_function_var.get()
            if func == "PDF转图片":
                self.convert_to_images()
            else:
                self.convert_to_word()
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("转换失败", f"转换过程中出错：\n{str(e)}"))
            self.root.after(0, lambda: self.status_message.set("转换失败"))
        finally:
            # 重新启用转换按钮
            self.conversion_active = False
            self.stop_page_timer()
            self.root.after(0, lambda: self.convert_btn.config(state=tk.NORMAL))
    
    def convert_to_images(self):
        """将PDF批量转换为图片"""
        import fitz as fitz_lib

        files = self.selected_files_list
        if not files:
            self.root.after(0, lambda: messagebox.showwarning("提示", "请先选择PDF文件！"))
            return

        # 解析DPI
        try:
            dpi = int(self.image_dpi_var.get())
            if dpi < 36 or dpi > 1200:
                raise ValueError
        except ValueError:
            self.root.after(0, lambda: messagebox.showerror("参数错误", "DPI必须是36-1200之间的整数"))
            return

        img_format = self.image_format_var.get().upper()
        if img_format not in ("PNG", "JPEG"):
            img_format = "PNG"
        ext = ".png" if img_format == "PNG" else ".jpg"
        zoom = dpi / 72.0

        # 计算总页数（用于进度）
        total_pages_all = 0
        file_page_counts = []
        for f in files:
            try:
                doc = fitz_lib.open(f)
                count = len(doc)
                doc.close()
                file_page_counts.append(count)
                total_pages_all += count
            except Exception as e:
                self.root.after(0, lambda msg=str(e), fn=f: messagebox.showerror(
                    "文件错误", f"无法打开: {os.path.basename(fn)}\n{msg}"))
                return

        if total_pages_all == 0:
            self.root.after(0, lambda: messagebox.showwarning("提示", "所有PDF文件均无内容"))
            return

        # 获取页范围
        page_start_text = self.page_start_var.get().strip()
        page_end_text = self.page_end_var.get().strip()
        use_range = bool(page_start_text or page_end_text)

        self.root.after(0, lambda: self.progress_bar.config(mode='determinate', maximum=100, value=0))
        self.start_time = time.time()
        processed = 0
        output_dirs = []
        errors = []

        for file_idx, pdf_path in enumerate(files):
            basename = os.path.splitext(os.path.basename(pdf_path))[0]
            output_dir = os.path.join(os.path.dirname(pdf_path), basename)

            # 如果文件夹已存在，加时间戳
            if os.path.exists(output_dir):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_dir = os.path.join(os.path.dirname(pdf_path), f"{basename}_{timestamp}")
            os.makedirs(output_dir, exist_ok=True)
            output_dirs.append(output_dir)

            try:
                doc = fitz_lib.open(pdf_path)
                page_count = len(doc)

                # 确定页范围
                start_idx = 0
                end_idx = page_count
                if use_range:
                    try:
                        s = int(page_start_text) if page_start_text else 1
                        e = int(page_end_text) if page_end_text else page_count
                        s = max(1, min(s, page_count))
                        e = max(s, min(e, page_count))
                        start_idx = s - 1
                        end_idx = e
                    except ValueError:
                        pass  # 忽略无效范围，处理全部页

                file_label = os.path.basename(pdf_path)
                for page_idx in range(start_idx, end_idx):
                    page = doc[page_idx]
                    mat = fitz_lib.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat, alpha=False)

                    img_filename = f"{page_idx}{ext}"
                    img_path = os.path.join(output_dir, img_filename)

                    if img_format == "JPEG":
                        pix.save(img_path, jpg_quality=95)
                    else:
                        pix.save(img_path)

                    processed += 1
                    progress = int(processed / total_pages_all * 100)
                    page_num = page_idx + 1
                    self.root.after(0, lambda v=progress: self.progress_bar.config(value=v))
                    status = f"[{file_idx+1}/{len(files)}] {file_label} - 第{page_num}页 ({progress}%)"
                    self.root.after(0, lambda t=status: self.set_progress_text(t))
                    self.base_status_text = f"正在转换: {file_label}"
                    self.root.after(0, self.apply_status_text)

                doc.close()
            except Exception as e:
                errors.append(f"{os.path.basename(pdf_path)}: {str(e)}")
                logging.error(f"PDF转图片失败 [{pdf_path}]: {e}")

        # 完成
        self.root.after(0, lambda: self.progress_bar.config(value=100))
        self.root.after(0, lambda: self.set_progress_text("转换完成！(100%)"))

        if errors:
            err_msg = "\n".join(errors)
            success_msg = f"转换完成，但有 {len(errors)} 个文件出错：\n\n{err_msg}"
            if output_dirs:
                success_msg += f"\n\n成功的文件已保存到各PDF同目录下的文件夹中"
            self.root.after(0, lambda: messagebox.showwarning("部分完成", success_msg))
        else:
            if len(files) == 1:
                success_msg = f"PDF已成功转换为图片！\n\nDPI: {dpi}  格式: {img_format}\n共 {processed} 页\n\n保存位置：\n{output_dirs[0]}"
            else:
                dir_list = "\n".join(output_dirs[:5])
                if len(output_dirs) > 5:
                    dir_list += f"\n...等共 {len(output_dirs)} 个文件夹"
                success_msg = f"所有PDF已成功转换为图片！\n\nDPI: {dpi}  格式: {img_format}\n共 {len(files)} 个文件，{processed} 页\n\n保存位置：\n{dir_list}"
            self.root.after(0, lambda: messagebox.showinfo("转换成功", success_msg))

        # 打开第一个输出文件夹
        if output_dirs:
            self.root.after(0, lambda d=output_dirs[0]: self.open_folder_direct(d))

        self.root.after(0, lambda: self.status_message.set(
            f"转换完成：{len(files)}个文件，共{processed}页"))

    def open_folder_direct(self, folder_path):
        """直接打开文件夹"""
        try:
            os.startfile(folder_path)
        except Exception as e:
            logging.warning(f"无法打开文件夹: {e}")

    def convert_to_word(self):
        """将PDF转换为Word"""
        if not PDF2DOCX_AVAILABLE or ProgressConverter is None:
            self.root.after(0, lambda: messagebox.showerror("错误", "pdf2docx库未安装！\n请运行: pip install pdf2docx"))
            return

        ocr_on = self.ocr_enabled_var.get()
        formula_api_on = self.formula_api_enabled_var.get()

        # 检查API配置
        if (ocr_on or formula_api_on) and (not self.baidu_api_key or not self.baidu_secret_key):
            self.root.after(0, lambda: messagebox.showerror(
                "API未配置", "您启用了OCR或公式识别功能，但尚未配置百度API。\n请点击设置按钮 ⚙ 配置API Key。"))
            return

        # 更新状态
        self.base_status_text = "正在初始化转换..."
        self.root.after(0, self.apply_status_text)
        self.root.after(0, lambda: self.set_progress_text("准备中..."))

        # 生成输出文件名
        input_file = self.selected_file.get()
        output_file = self.generate_output_filename(input_file, '.docx')

        # 执行转换
        self.root.after(0, lambda: self.progress_bar.config(mode='determinate', maximum=100, value=0))

        try:
            if ocr_on:
                # OCR模式：渲染PDF为图片 → 百度OCR识别 → 生成Word
                self._convert_with_ocr(input_file, output_file, formula_api_on)
            else:
                # 普通模式：pdf2docx转换 + 可选的公式API后处理
                self._convert_with_pdf2docx(input_file, output_file, formula_api_on)

        except Exception as e:
            raise e

    def _convert_with_pdf2docx(self, input_file, output_file, formula_api_on):
        """使用pdf2docx转换，可选公式API后处理"""
        formula_mode = 0  # 始终用普通模式
        cv = ProgressConverter(input_file, progress_callback=self.update_progress,
                               formula_mode=formula_mode)
        self.total_pages = len(cv.fitz_doc)
        if self.total_pages <= 0:
            raise ConversionException("无法读取PDF页数")
        try:
            start_idx, end_idx, range_total = self.get_page_range(self.total_pages)
        except ValueError as e:
            self.root.after(0, lambda: messagebox.showerror("页范围错误", str(e)))
            self.root.after(0, lambda: self.status_message.set("页范围无效"))
            cv.close()
            return

        self.total_steps = range_total * 2
        self.start_time = time.time()
        self.root.after(0, lambda: self.set_progress_text(f"共 {range_total} 页，开始转换..."))
        cv.convert(output_file, start=start_idx, end=end_idx)
        cv.close()

        # 公式API后处理
        formula_fix_count = 0
        if formula_api_on:
            actual_end = end_idx if end_idx is not None else self.total_pages
            self.root.after(0, lambda: self.set_progress_text("正在检测公式页面..."))
            math_pages = detect_math_pages(fitz.open(input_file), start=start_idx, end=actual_end)
            if math_pages:
                self.root.after(0, lambda: self.set_progress_text("正在调用API识别公式..."))
                formula_fix_count = self._post_process_formula_api(
                    output_file, input_file, math_pages)

        # 转换成功
        self.root.after(0, lambda: self.progress_bar.config(value=100))
        self.root.after(0, lambda: self.set_progress_text("转换完成！(100%)"))

        success_msg = f"PDF已成功转换为Word！\n\n保存位置：\n{output_file}"
        if formula_fix_count > 0:
            success_msg += f"\n\n已识别并转换 {formula_fix_count} 处数学公式为可编辑格式"
        success_msg += "\n\n是否打开文件所在文件夹？"

        if messagebox.askyesno("转换成功", success_msg):
            self.open_folder(output_file)

        if cv.skipped_pages:
            skipped_text = self.format_skipped_pages(cv.skipped_pages)
            messagebox.showwarning("跳过异常页", f"以下页面在转换中被跳过：\n{skipped_text}")

    def _convert_with_ocr(self, input_file, output_file, formula_api_on):
        """OCR模式：整页渲染为图片 → 百度OCR识别文字 → 生成Word"""
        fitz_doc = fitz.open(input_file)
        self.total_pages = len(fitz_doc)
        if self.total_pages <= 0:
            raise RuntimeError("无法读取PDF页数")
        try:
            start_idx, end_idx, range_total = self.get_page_range(self.total_pages)
        except ValueError as e:
            self.root.after(0, lambda: messagebox.showerror("页范围错误", str(e)))
            self.root.after(0, lambda: self.status_message.set("页范围无效"))
            fitz_doc.close()
            return

        actual_end = end_idx if end_idx is not None else self.total_pages
        self.start_time = time.time()
        client = self._get_baidu_client()
        doc = Document()
        formula_count = 0
        ocr_errors = []

        for i, page_idx in enumerate(range(start_idx, actual_end)):
            page_num = page_idx + 1
            percent = int(((i + 0.5) / range_total) * 100)
            self.root.after(0, lambda p=percent, pn=page_num: (
                self.progress_bar.config(value=p),
                self.set_progress_text(f"OCR识别第 {pn} 页... ({p}%)")
            ))
            self.base_status_text = f"正在OCR识别第 {page_num} 页，共 {range_total} 页"
            self.root.after(0, self.apply_status_text)

            pdf_page = fitz_doc[page_idx]
            dpi = 300
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = pdf_page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # 添加分页符
            if i > 0:
                run = doc.add_paragraph().add_run()
                run.add_break(WD_BREAK.PAGE)

            # ---- 文字OCR识别 ----
            text_lines = None
            try:
                text_lines = client.recognize_text(img_bytes)
                logging.info(f"Page {page_num}: OCR recognized {len(text_lines)} lines")
            except Exception as e:
                err_msg = f"第{page_num}页OCR失败: {e}"
                logging.error(err_msg)
                ocr_errors.append(err_msg)

            if text_lines:
                for line_text in text_lines:
                    doc.add_paragraph(line_text)
            else:
                # OCR失败或无内容，用图片替代
                logging.info(f"Page {page_num}: No text recognized, inserting image")
                img_stream = io.BytesIO(img_bytes)
                page_width = pdf_page.rect.width / 72.0
                doc.add_picture(img_stream, width=Inches(min(page_width, 6.3)))

            # ---- 公式识别（每页都尝试，不依赖字体检测）----
            if formula_api_on:
                try:
                    self.root.after(0, lambda pn=page_num: self.set_progress_text(
                        f"识别第 {pn} 页公式..."))
                    formulas = client.recognize_formula(img_bytes)
                    for latex_str in formulas:
                        if not latex_str.strip():
                            continue
                        omml_elem = latex_to_omml(latex_str, self.xslt_path)
                        if omml_elem is not None:
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            insert_omml_to_paragraph(para, omml_elem)
                            formula_count += 1
                        else:
                            # OMML转换失败，以LaTeX文本形式插入
                            para = doc.add_paragraph(f"[公式] {latex_str}")
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            formula_count += 1
                except Exception as e:
                    logging.warning(f"Page {page_num} formula API error: {e}")

        doc.save(output_file)
        fitz_doc.close()

        # 转换成功
        self.root.after(0, lambda: self.progress_bar.config(value=100))
        self.root.after(0, lambda: self.set_progress_text("转换完成！(100%)"))

        success_msg = f"PDF已成功转换为Word（OCR模式）！\n\n保存位置：\n{output_file}"
        success_msg += f"\n\n共处理 {range_total} 页"
        if formula_count > 0:
            success_msg += f"，识别 {formula_count} 处公式"
        if ocr_errors:
            success_msg += f"\n\n⚠ {len(ocr_errors)} 页识别出错（已用图片替代）"
        success_msg += "\n\n是否打开文件所在文件夹？"

        if messagebox.askyesno("转换成功", success_msg):
            self.open_folder(output_file)

        # 显示错误详情
        if ocr_errors:
            err_detail = "\n".join(ocr_errors[:10])
            messagebox.showwarning("OCR识别警告", f"以下页面识别失败：\n{err_detail}")

    def _post_process_formula_api(self, docx_path, pdf_path, math_page_ids):
        """使用百度API识别公式并替换为Word原生OMML公式"""
        doc_obj = Document(docx_path)
        fitz_doc = fitz.open(pdf_path)
        client = self._get_baidu_client()
        fix_count = 0

        # 第一步：规范化简单的数学Unicode字符
        for para in doc_obj.paragraphs:
            for run in para.runs:
                if _has_math_unicode(run.text):
                    run.text = _normalize_math_unicode(run.text)
                    fix_count += 1
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if _has_math_unicode(run.text):
                                run.text = _normalize_math_unicode(run.text)
                                fix_count += 1

        # 第二步：找到独立公式块，裁剪发送API识别
        for page_id in sorted(math_page_ids):
            pdf_page = fitz_doc[page_id]
            td = pdf_page.get_text("dict")
            for block in td.get("blocks", []):
                if not _is_display_equation(block):
                    continue

                # 裁剪公式区域
                bbox = block["bbox"]
                x0, y0, x1, y1 = bbox
                padding = 5
                clip = fitz.Rect(
                    max(0, x0 - padding), max(0, y0 - padding),
                    min(pdf_page.rect.width, x1 + padding),
                    min(pdf_page.rect.height, y1 + padding)
                )
                if clip.is_empty or clip.width < 5 or clip.height < 5:
                    continue

                dpi = 300
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = pdf_page.get_pixmap(matrix=mat, clip=clip)
                img_bytes = pix.tobytes("png")

                # 调用公式API
                try:
                    self.root.after(0, lambda pid=page_id: self.set_progress_text(
                        f"正在识别第 {pid + 1} 页的公式..."))
                    formulas = client.recognize_formula(img_bytes)
                except Exception as e:
                    logging.warning(f"Formula API error on page {page_id + 1}: {e}")
                    continue

                if not formulas:
                    continue

                latex_str = formulas[0]  # 取第一个结果

                # 在docx中找到对应的段落并替换
                block_text = _get_block_text(block)
                norm_text = _normalize_math_unicode(block_text)
                norm_compact = ''.join(norm_text.split())
                if len(norm_compact) < 2:
                    continue

                for para in doc_obj.paragraphs:
                    para_compact = ''.join(para.text.split())
                    if len(para_compact) < 2:
                        continue
                    if self._text_similar(para_compact, norm_compact):
                        # 尝试转为OMML
                        omml_elem = latex_to_omml(latex_str, self.xslt_path)
                        if omml_elem is not None:
                            # 清除原有内容
                            for run in para.runs:
                                run.text = ""
                            # 清除段落中残余的 XML 子元素
                            for child in list(para._element):
                                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                                if tag == 'r':
                                    para._element.remove(child)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            insert_omml_to_paragraph(para, omml_elem)
                            fix_count += 1
                            logging.info(f"Replaced equation with OMML: {latex_str[:50]}")
                        else:
                            # OMML失败，用裁剪图片替代
                            for run in para.runs:
                                run.text = ""
                            img_stream = io.BytesIO(img_bytes)
                            region_width = clip.width / 72.0
                            doc_obj.add_picture(img_stream, width=Inches(min(region_width, 6.0)))
                            body = doc_obj.element.body
                            pic_element = body[-1]
                            para._element.addnext(pic_element)
                            fix_count += 1
                        break

        doc_obj.save(docx_path)
        fitz_doc.close()
        return fix_count

    @staticmethod
    def _text_similar(a, b):
        """判断两个文本（已去空白）是否相似"""
        if not a or not b:
            return False
        if a == b:
            return True
        shorter = min(len(a), len(b))
        longer = max(len(a), len(b))
        if shorter < 3 or shorter / longer < 0.3:
            return False
        # 计算公共字符比例
        set_a, set_b = set(a), set(b)
        common_chars = set_a & set_b
        all_chars = set_a | set_b
        if not all_chars:
            return False
        jaccard = len(common_chars) / len(all_chars)
        # 检查子串包含
        if shorter >= 4 and (a[:shorter] in b or b[:shorter] in a):
            return True
        return jaccard > 0.6

    @staticmethod
    def _replace_para_with_equation_image(doc_obj, para, pdf_page, bbox):
        """将段落内容替换为PDF裁剪的公式图片"""
        x0, y0, x1, y1 = bbox
        padding = 2
        x0 = max(0, x0 - padding)
        y0 = max(0, y0 - padding)
        x1 = min(pdf_page.rect.width, x1 + padding)
        y1 = min(pdf_page.rect.height, y1 + padding)

        clip = fitz.Rect(x0, y0, x1, y1)
        if clip.is_empty or clip.width < 1 or clip.height < 1:
            return

        dpi = 300
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = pdf_page.get_pixmap(matrix=mat, clip=clip)
        img_bytes = pix.tobytes("png")

        region_width = clip.width / 72.0
        target_width = min(region_width, 6.0)

        # 清除段落中的所有Run的文字
        for run in para.runs:
            run.text = ""

        # 通过Document.add_picture添加图片（会追加到末尾），然后移动到当前段落后面
        img_stream = io.BytesIO(img_bytes)
        doc_obj.add_picture(img_stream, width=Inches(target_width))
        # 获取新创建的图片段落（body的最后一个子元素）
        body = doc_obj.element.body
        pic_element = body[-1]
        # 将图片段落移到当前段落的后面
        para._element.addnext(pic_element)

    def update_progress(self, phase: str, current: int, total: int, page_id: int):
        """更新进度条和提示信息"""
        if total <= 0:
            return

        total_steps = total * 2
        if phase in ('start-parse', 'start-make'):
            phase_text = "解析" if phase == 'start-parse' else "生成"
            self.current_phase = phase_text
            self.current_page_id = page_id
            self.current_page_index = current
            self.current_page_total = total
            self.page_start_time = time.time()
            self.base_status_text = f"正在{phase_text}第 {page_id} 页，共 {total} 页"
            self.root.after(0, self.apply_status_text)
            return

        if phase in ('skip-parse', 'skip-make'):
            phase_text = "解析" if phase == 'skip-parse' else "生成"
            self.base_status_text = f"第 {page_id} 页{phase_text}失败，已跳过"
            self.root.after(0, self.apply_status_text)
            return

        if phase == 'parse':
            completed_steps = current
            percent = int(round((completed_steps / total_steps) * 100))
            phase_text = "解析"
        else:
            completed_steps = total + current
            percent = int(round((completed_steps / total_steps) * 100))
            phase_text = "生成"

        page_text = self.format_page_text(phase_text, current, total, page_id)
        self.base_status_text = f"正在{phase_text}第 {page_id} 页，共 {total} 页"

        eta_text = ""
        if self.start_time and completed_steps > 0:
            elapsed = time.time() - self.start_time
            remaining_steps = max(total_steps - completed_steps, 0)
            eta_seconds = int(round(elapsed * remaining_steps / completed_steps))
            eta_text = f"，预计剩余 {self.format_eta(eta_seconds)}"
        self.current_eta_text = eta_text

        def _apply():
            self.progress_bar.config(mode='determinate', maximum=100)
            self.progress_bar['value'] = percent
            self.set_progress_text(f"{page_text} ({percent}%)")
            self.apply_status_text()

        self.root.after(0, _apply)

    @staticmethod
    def format_eta(seconds: int) -> str:
        """格式化预计剩余时间"""
        minutes, sec = divmod(max(seconds, 0), 60)
        hours, minutes = divmod(minutes, 60)
        if hours > 0:
            return f"{hours}小时{minutes}分{sec}秒"
        if minutes > 0:
            return f"{minutes}分{sec}秒"
        return f"{sec}秒"

    def start_page_timer(self):
        if self.page_timer_job is not None:
            return
        self.page_timer_job = self.root.after(1000, self.refresh_page_timer)

    def stop_page_timer(self):
        if self.page_timer_job is not None:
            try:
                self.root.after_cancel(self.page_timer_job)
            except Exception:
                pass
            self.page_timer_job = None

    def refresh_page_timer(self):
        self.apply_status_text()
        if self.conversion_active:
            self.page_timer_job = self.root.after(1000, self.refresh_page_timer)
        else:
            self.page_timer_job = None

    def apply_status_text(self):
        text = self.base_status_text or ""
        if self.current_eta_text:
            text += self.current_eta_text
        if self.page_start_time:
            elapsed = int(time.time() - self.page_start_time)
            text += f"，当前页耗时 {self.format_eta(elapsed)}"
            if elapsed >= self.page_timeout_seconds:
                text += "，该页复杂请耐心等待"
        if text:
            self.status_message.set(text)

    def format_page_text(self, phase_text: str, current: int, total: int, page_id: int) -> str:
        if self.total_pages and total != self.total_pages:
            return f"{phase_text}页 {current}/{total} (原页 {page_id})"
        return f"{phase_text}页 {page_id}/{total}"

    def open_settings_window(self):
        """打开设置窗口（含API配置）"""
        win = tk.Toplevel(self.root)
        win.title("设置")
        win.geometry("480x520")
        win.resizable(False, False)

        # 使用 Notebook 分页签
        notebook = ttk.Notebook(win)
        notebook.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)

        # ========== 页签1：外观设置 ==========
        tab_appearance = tk.Frame(notebook, padx=12, pady=12)
        notebook.add(tab_appearance, text="外观设置")

        tk.Label(tab_appearance, text="标题文字:", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)
        title_entry = tk.Entry(tab_appearance, textvariable=self.title_text_var,
                               font=("Microsoft YaHei", 10))
        title_entry.pack(fill=tk.X, pady=(4, 12))

        tk.Button(tab_appearance, text="更换背景", font=("Microsoft YaHei", 10),
                  command=self.choose_background_image).pack(anchor=tk.W)

        tk.Label(tab_appearance, text="面板透明度:", font=("Microsoft YaHei", 10)
                 ).pack(anchor=tk.W, pady=(12, 0))
        tk.Scale(tab_appearance, from_=0, to=100, orient=tk.HORIZONTAL,
                 resolution=1, showvalue=True, variable=self.panel_opacity_var,
                 command=self.on_opacity_change).pack(fill=tk.X, pady=(4, 0))

        tk.Button(tab_appearance, text="应用标题", font=("Microsoft YaHei", 10),
                  command=self.apply_title_text).pack(anchor=tk.W, pady=(12, 0))

        # ========== 页签2：API设置 ==========
        tab_api = tk.Frame(notebook, padx=12, pady=12)
        notebook.add(tab_api, text="API设置")

        # 百度OCR配置
        tk.Label(tab_api, text="百度OCR API（用于文字识别和公式识别）",
                 font=("Microsoft YaHei", 10, "bold")).pack(anchor=tk.W, pady=(0, 8))

        tk.Label(tab_api, text="API Key:", font=("Microsoft YaHei", 9)).pack(anchor=tk.W)
        api_key_var = tk.StringVar(value=self.baidu_api_key)
        tk.Entry(tab_api, textvariable=api_key_var, font=("Microsoft YaHei", 9),
                 width=50).pack(fill=tk.X, pady=(2, 6))

        tk.Label(tab_api, text="Secret Key:", font=("Microsoft YaHei", 9)).pack(anchor=tk.W)
        secret_key_var = tk.StringVar(value=self.baidu_secret_key)
        tk.Entry(tab_api, textvariable=secret_key_var, font=("Microsoft YaHei", 9),
                 width=50, show="*").pack(fill=tk.X, pady=(2, 8))

        # 测试连接
        test_status_var = tk.StringVar(value="")
        test_frame = tk.Frame(tab_api)
        test_frame.pack(fill=tk.X, pady=(0, 8))

        def do_test():
            ak = api_key_var.get().strip()
            sk = secret_key_var.get().strip()
            if not ak or not sk:
                test_status_var.set("⚠ 请填写API Key和Secret Key")
                return
            test_status_var.set("⏳ 正在测试...")
            win.update()
            client = BaiduOCRClient(ak, sk)
            ok, msg = client.test_connection()
            if ok:
                test_status_var.set("✅ 连接成功")
            else:
                test_status_var.set(f"❌ 失败: {msg[:50]}")

        tk.Button(test_frame, text="测试连接", font=("Microsoft YaHei", 9),
                  command=do_test).pack(side=tk.LEFT)
        tk.Label(test_frame, textvariable=test_status_var,
                 font=("Microsoft YaHei", 9)).pack(side=tk.LEFT, padx=(10, 0))

        # 说明
        hint_text = (
            "注册地址：https://cloud.baidu.com/product/ocr\n"
            "1. 注册百度智能云账号\n"
            "2. 创建文字识别应用，获取API Key和Secret Key\n"
            "3. 同一个应用可同时使用文字识别和公式识别\n"
            "4. 免费额度：通用文字500次/月"
        )
        tk.Label(tab_api, text=hint_text, font=("Microsoft YaHei", 8),
                 fg="#666666", justify=tk.LEFT, wraplength=420).pack(anchor=tk.W, pady=(4, 12))

        # XSLT路径（高级选项）
        tk.Label(tab_api, text="高级选项（通常无需修改）:",
                 font=("Microsoft YaHei", 8), fg="#aaaaaa").pack(anchor=tk.W, pady=(8, 0))
        xslt_hint = "留空自动检测Office安装路径，仅Office路径异常时手动填写"
        tk.Label(tab_api, text=f"MML2OMML.XSL: {xslt_hint}",
                 font=("Microsoft YaHei", 8), fg="#aaaaaa").pack(anchor=tk.W)
        xslt_var = tk.StringVar(value=self.xslt_path or "")
        tk.Entry(tab_api, textvariable=xslt_var, font=("Microsoft YaHei", 8),
                 fg="#aaaaaa").pack(fill=tk.X, pady=(2, 0))

        # 保存按钮
        def save_api_settings():
            self.baidu_api_key = api_key_var.get().strip()
            self.baidu_secret_key = secret_key_var.get().strip()
            self.xslt_path = xslt_var.get().strip() or None
            self._baidu_client = None  # 重建客户端
            self.save_settings()
            self._update_api_hint()
            messagebox.showinfo("设置", "API设置已保存", parent=win)

        tk.Button(tab_api, text="保存设置", font=("Microsoft YaHei", 10, "bold"),
                  command=save_api_settings).pack(anchor=tk.E, pady=(12, 0))

    def apply_title_text(self):
        text = self.title_text_var.get().strip() or "程新伟专属转换器"
        self.title_text_var.set(text)
        self.save_settings()

    def on_opacity_change(self, _value=None):
        self.apply_panel_image()
        self.save_settings()

    def _on_option_changed(self):
        """复选框状态变化时更新提示"""
        self._update_api_hint()
        self.save_settings()

    def _on_function_changed(self, event=None):
        """功能选择变化时切换选项区域"""
        func = self.current_function_var.get()
        if func == "PDF转Word":
            self.panel_canvas.itemconfigure(self.cv_formula_frame, state='normal')
            self.panel_canvas.itemconfigure(self.cv_api_hint, state='normal')
            self.panel_canvas.itemconfigure(self.cv_image_options, state='hidden')
            self.root.title("程新伟专属转换器 - PDF转Word")
        elif func == "PDF转图片":
            self.panel_canvas.itemconfigure(self.cv_formula_frame, state='hidden')
            self.panel_canvas.itemconfigure(self.cv_api_hint, state='hidden')
            self.panel_canvas.itemconfigure(self.cv_image_options, state='normal')
            self.root.title("程新伟专属转换器 - PDF转图片")
        # 切换功能时清除已选文件
        self.selected_file.set("")
        self.selected_files_list = []
        self.status_message.set("就绪")
        self.save_settings()

    def _update_api_hint(self):
        """更新API状态提示文字"""
        if not self.panel_canvas:
            return
        ocr_on = self.ocr_enabled_var.get()
        formula_on = self.formula_api_enabled_var.get()
        if not ocr_on and not formula_on:
            self.panel_canvas.itemconfigure(self.cv_api_hint, text="")
            return
        has_key = bool(self.baidu_api_key and self.baidu_secret_key)
        parts = []
        if ocr_on:
            parts.append("OCR识别")
        if formula_on:
            parts.append("公式识别")
        feature_text = " + ".join(parts)
        if has_key:
            self.panel_canvas.itemconfigure(
                self.cv_api_hint,
                text=f"已启用: {feature_text}（百度API已配置）",
                fill="#228B22"
            )
        else:
            self.panel_canvas.itemconfigure(
                self.cv_api_hint,
                text=f"已启用: {feature_text}（⚠ 请在设置中配置API Key）",
                fill="#CC0000"
            )

    def _get_baidu_client(self):
        """获取或创建百度OCR客户端"""
        if not REQUESTS_AVAILABLE:
            raise RuntimeError("requests库未安装，请运行: pip install requests")
        if not self.baidu_api_key or not self.baidu_secret_key:
            raise RuntimeError("百度OCR API未配置，请在设置中填写API Key和Secret Key")
        if self._baidu_client is None:
            self._baidu_client = BaiduOCRClient(self.baidu_api_key, self.baidu_secret_key)
        return self._baidu_client

    def choose_background_image(self):
        filename = filedialog.askopenfilename(
            title="选择背景图片",
            filetypes=[("图片文件", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"), ("所有文件", "*.*")]
        )
        if not filename:
            return

        if not PIL_AVAILABLE:
            messagebox.showerror("错误", "Pillow库未安装，无法加载图片背景。\n请运行: pip install Pillow")
            return

        try:
            app_dir = self.get_app_dir()
            ext = os.path.splitext(filename)[1].lower() or ".png"
            target = os.path.join(app_dir, f"background{ext}")
            shutil.copyfile(filename, target)
            self.bg_image_path = target
            self.apply_background_image()
            self.save_settings()
        except Exception as e:
            messagebox.showerror("错误", f"无法设置背景图片：\n{str(e)}")

    def apply_background_image(self):
        if not PIL_AVAILABLE:
            return
        if not self.bg_image_path or not os.path.exists(self.bg_image_path):
            return

        try:
            img = Image.open(self.bg_image_path)
            width = self.root.winfo_width()
            height = self.root.winfo_height()
            if width <= 1 or height <= 1:
                self.root.update_idletasks()
                width = self.root.winfo_width()
                height = self.root.winfo_height()
            img = img.resize((width, height), Image.LANCZOS).convert("RGB")
            self.bg_pil = img
            self.bg_image = ImageTk.PhotoImage(img)
            if self.bg_label is None:
                self.bg_label = tk.Label(self.root, image=self.bg_image)
                self.bg_label.place(x=0, y=0, relwidth=1, relheight=1)
                self.bg_label.lower()
            else:
                self.bg_label.configure(image=self.bg_image)

            self.root.after(0, self.apply_panel_image)
        except Exception as e:
            messagebox.showerror("错误", f"背景图片加载失败：\n{str(e)}")

    def on_root_resize(self, event):
        if not self.bg_image_path:
            return
        if self.resize_job is not None:
            try:
                self.root.after_cancel(self.resize_job)
            except Exception:
                pass
        self.resize_job = self.root.after(200, self.apply_background_image)

    def on_panel_resize(self, event):
        self.layout_canvas()
        if self.panel_resize_job is not None:
            try:
                self.root.after_cancel(self.panel_resize_job)
            except Exception:
                pass
        self.panel_resize_job = self.root.after(50, self.apply_panel_image)

    def refresh_layout(self):
        self.root.update_idletasks()
        self.layout_canvas()
        self.apply_panel_image()

    def apply_panel_image(self):
        if not PIL_AVAILABLE:
            return
        if not self.bg_pil or self.panel_canvas is None:
            return

        width = self.root.winfo_width()
        height = self.root.winfo_height()
        panel_width = max(width - self.panel_padding * 2, 1)
        panel_height = max(height - self.panel_padding * 2, 1)

        if self.bg_pil.size[0] != width or self.bg_pil.size[1] != height:
            return

        left = self.panel_padding
        top = self.panel_padding
        right = left + panel_width
        bottom = top + panel_height

        panel_img = self.bg_pil.crop((left, top, right, bottom))
        opacity = max(0.2, min(1.0, self.panel_opacity_var.get() / 100.0))
        overlay = Image.new("RGB", panel_img.size, (255, 255, 255))
        panel_img = Image.blend(overlay, panel_img, opacity)
        self.panel_image = ImageTk.PhotoImage(panel_img)
        if self.panel_image_id is None:
            self.panel_image_id = self.panel_canvas.create_image(0, 0, anchor="nw", image=self.panel_image)
            self.panel_canvas.tag_lower(self.panel_image_id)
        else:
            self.panel_canvas.itemconfigure(self.panel_image_id, image=self.panel_image)
        self.panel_canvas.update_idletasks()

    def load_settings(self):
        if not os.path.exists(self.settings_path):
            return
        try:
            with open(self.settings_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            title_text = data.get('title_text')
            if title_text:
                self.title_text_var.set(title_text)
            bg_path = data.get('background_image')
            if bg_path and os.path.exists(bg_path):
                self.bg_image_path = bg_path
            opacity = data.get('panel_opacity', data.get('background_opacity'))
            if isinstance(opacity, (int, float)):
                self.panel_opacity_var.set(max(20.0, min(100.0, float(opacity))))
            # 兼容旧版 formula_mode
            formula_mode = data.get('formula_mode', 0)
            if isinstance(formula_mode, bool):
                formula_mode = 1 if formula_mode else 0
            self.formula_mode_var.set(int(formula_mode))
            # 新增：OCR和公式选项
            self.ocr_enabled_var.set(data.get('ocr_enabled', False))
            self.formula_api_enabled_var.set(data.get('formula_api_enabled', False))
            # API 配置
            self.baidu_api_key = _simple_decrypt(data.get('baidu_api_key_enc', ''))
            self.baidu_secret_key = _simple_decrypt(data.get('baidu_secret_key_enc', ''))
            self.xslt_path = data.get('xslt_path') or None
            # 功能选择和图片选项
            saved_func = data.get('current_function', 'PDF转Word')
            if saved_func in ("PDF转Word", "PDF转图片"):
                self.current_function_var.set(saved_func)
                self._on_function_changed()
            saved_dpi = data.get('image_dpi', '200')
            if saved_dpi:
                self.image_dpi_var.set(str(saved_dpi))
            saved_fmt = data.get('image_format', 'PNG')
            if saved_fmt in ('PNG', 'JPEG'):
                self.image_format_var.set(saved_fmt)
            if self.bg_image_path:
                self.apply_background_image()
            self._update_api_hint()
        except Exception:
            pass

    def save_settings(self):
        data = {
            'title_text': self.title_text_var.get().strip(),
            'background_image': self.bg_image_path,
            'panel_opacity': float(self.panel_opacity_var.get()),
            'formula_mode': int(self.formula_mode_var.get()),
            'ocr_enabled': bool(self.ocr_enabled_var.get()),
            'formula_api_enabled': bool(self.formula_api_enabled_var.get()),
            'baidu_api_key_enc': _simple_encrypt(self.baidu_api_key),
            'baidu_secret_key_enc': _simple_encrypt(self.baidu_secret_key),
            'xslt_path': self.xslt_path or '',
            'current_function': self.current_function_var.get(),
            'image_dpi': self.image_dpi_var.get(),
            'image_format': self.image_format_var.get(),
        }
        try:
            with open(self.settings_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    @staticmethod
    def get_app_dir():
        if getattr(sys, 'frozen', False):
            return os.path.dirname(sys.executable)
        return os.path.dirname(os.path.abspath(__file__))

    def get_page_range(self, total_pages: int):
        start_text = self.page_start_var.get().strip()
        end_text = self.page_end_var.get().strip()

        if not start_text and not end_text:
            return 0, None, total_pages

        if start_text and not start_text.isdigit():
            raise ValueError("起始页必须是数字")
        if end_text and not end_text.isdigit():
            raise ValueError("结束页必须是数字")

        start_page = int(start_text) if start_text else 1
        end_page = int(end_text) if end_text else total_pages

        if start_page < 1 or end_page < 1:
            raise ValueError("页码必须从1开始")
        if start_page > end_page:
            raise ValueError("起始页不能大于结束页")
        if end_page > total_pages:
            raise ValueError("结束页超出总页数")

        start_idx = start_page - 1
        end_idx = end_page
        return start_idx, end_idx, end_page - start_idx

    @staticmethod
    def format_skipped_pages(skipped_pages):
        pages = sorted(set(skipped_pages))
        if len(pages) <= 30:
            return ", ".join(str(p) for p in pages)
        head = ", ".join(str(p) for p in pages[:30])
        return f"{head} ...（共 {len(pages)} 页）"
    
    def generate_output_filename(self, input_file, extension):
        """生成输出文件名"""
        # 获取输入文件的目录和基本名称
        directory = os.path.dirname(input_file)
        basename = os.path.splitext(os.path.basename(input_file))[0]
        
        # 添加时间戳避免覆盖
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{basename}_converted_{timestamp}{extension}"
        
        return os.path.join(directory, output_filename)
    
    def open_folder(self, filepath):
        """打开文件所在文件夹"""
        try:
            folder = os.path.dirname(os.path.abspath(filepath))
            os.startfile(folder)
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文件夹：\n{str(e)}")


def main():
    """主函数"""
    root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
