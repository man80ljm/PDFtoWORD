"""
PDF → Word 转换器

支持三种模式：
1. 普通模式：pdf2docx 直接转换
2. OCR模式：渲染→百度OCR→生成Word
3. 公式后处理：API识别公式并替换为OMML

通过 on_progress / pdf2docx_progress 回调报告进度，不直接操作UI。
"""

import io
import logging
import time

try:
    from pdf2docx.converter import ConversionException
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    import fitz
    _LIBS_AVAILABLE = True
except ImportError:
    _LIBS_AVAILABLE = False
    ConversionException = Exception

from core.math_utils import (
    detect_math_pages, has_math_unicode, normalize_math_unicode,
    is_display_equation, get_block_text,
    latex_to_omml, insert_omml_to_paragraph,
)
from core.ocr_client import BaiduOCRClient, REQUESTS_AVAILABLE
from core.progress_converter import ProgressConverter, PDF2DOCX_AVAILABLE


class PDFToWordConverter:
    """PDF→Word 转换器，与 UI 完全解耦。

    用法::

        converter = PDFToWordConverter(
            on_progress=my_callback,
            pdf2docx_progress=my_pdf2docx_callback,
        )
        result = converter.convert(input_file, output_file, ...)
    """

    def __init__(self, on_progress=None, pdf2docx_progress=None):
        """
        Args:
            on_progress: fn(percent, progress_text, status_text)
                简单进度回调。percent=-1 表示仅更新状态文字。
            pdf2docx_progress: fn(phase, current, total, page_id)
                用于 ProgressConverter 的详细阶段回调（解析/生成/跳过等）。
        """
        self.on_progress = on_progress or (lambda *a: None)
        self.pdf2docx_progress = pdf2docx_progress or (lambda *a: None)

    def _report(self, percent=-1, progress_text="", status_text=""):
        self.on_progress(percent, progress_text, status_text)

    # ----------------------------------------------------------
    # 公开入口
    # ----------------------------------------------------------

    def convert(self, input_file, output_file,
                start_page=0, end_page=None,
                ocr_enabled=False, formula_api_enabled=False,
                api_key=None, secret_key=None, xslt_path=None,
                ocr_mode="平衡"):
        """执行转换，返回结果字典。

        Returns:
            dict with keys:
                success (bool), message (str), output_file (str),
                formula_count (int), skipped_pages (set),
                errors (list[str]), page_count (int), mode (str)
        """
        result = {
            'success': False, 'message': '', 'output_file': output_file,
            'formula_count': 0, 'skipped_pages': set(),
            'errors': [], 'page_count': 0, 'mode': 'normal',
        }

        if not PDF2DOCX_AVAILABLE and not ocr_enabled:
            result['message'] = "pdf2docx库未安装！\n请运行: pip install pdf2docx"
            return result

        if (ocr_enabled or formula_api_enabled) and (not api_key or not secret_key):
            result['message'] = "您启用了OCR或公式识别功能，但尚未配置百度API。\n请点击设置按钮 ⚙ 配置API Key。"
            return result

        self._report(progress_text="准备中...", status_text="正在初始化转换...")

        try:
            if ocr_enabled:
                result['mode'] = 'ocr'
                self._convert_with_ocr(
                    input_file, output_file, start_page, end_page,
                    formula_api_enabled, api_key, secret_key, xslt_path, result,
                    ocr_mode=ocr_mode)
            else:
                result['mode'] = 'normal'
                self._convert_with_pdf2docx(
                    input_file, output_file, start_page, end_page,
                    formula_api_enabled, api_key, secret_key, xslt_path, result)
            result['success'] = True
        except Exception as e:
            result['message'] = str(e)
            logging.error(f"PDF→Word转换失败: {e}")

        return result

    # ----------------------------------------------------------
    # 普通模式（pdf2docx）
    # ----------------------------------------------------------

    def _convert_with_pdf2docx(self, input_file, output_file,
                               start_page, end_page,
                               formula_api_on, api_key, secret_key,
                               xslt_path, result):
        cv = ProgressConverter(input_file, progress_callback=self.pdf2docx_progress,
                               formula_mode=0)
        total_pages = len(cv.fitz_doc)
        if total_pages <= 0:
            raise ConversionException("无法读取PDF页数")

        result['page_count'] = total_pages
        self._report(progress_text=f"共 {total_pages} 页，开始转换...")

        end_idx = end_page if end_page else None
        cv.convert(output_file, start=start_page, end=end_idx)
        result['skipped_pages'] = cv.skipped_pages
        cv.close()

        # 公式API后处理
        if formula_api_on and api_key and secret_key:
            actual_end = end_page if end_page else total_pages
            self._report(progress_text="正在检测公式页面...")
            math_doc = fitz.open(input_file)
            math_pages = detect_math_pages(math_doc,
                                           start=start_page, end=actual_end)
            math_doc.close()
            if math_pages:
                self._report(progress_text="正在调用API识别公式...")
                client = BaiduOCRClient(api_key, secret_key)
                formula_count = self._post_process_formula_api(
                    output_file, input_file, math_pages, client, xslt_path)
                result['formula_count'] = formula_count

        self._report(percent=100, progress_text="转换完成！(100%)")

    # ----------------------------------------------------------
    # OCR模式
    # ----------------------------------------------------------

    def _convert_with_ocr(self, input_file, output_file,
                          start_page, end_page,
                          formula_api_on, api_key, secret_key,
                          xslt_path, result, ocr_mode="平衡"):
        fitz_doc = fitz.open(input_file)
        total_pages = len(fitz_doc)
        if total_pages <= 0:
            raise RuntimeError("无法读取PDF页数")

        actual_end = end_page if end_page else total_pages
        range_total = actual_end - start_page
        result['page_count'] = range_total

        client = BaiduOCRClient(api_key, secret_key)
        doc = Document()
        formula_count = 0
        ocr_errors = []
        dpi = self._ocr_mode_to_dpi(ocr_mode)

        for i, page_idx in enumerate(range(start_page, actual_end)):
            page_num = page_idx + 1
            percent = int(((i + 0.5) / range_total) * 100)
            self._report(
                percent=percent,
                progress_text=f"OCR识别第 {page_num} 页... ({percent}%)",
                status_text=f"正在OCR识别第 {page_num} 页，共 {range_total} 页",
            )

            # API调用频率控制，避免触发QPS限制
            if i > 0:
                time.sleep(0.5)

            pdf_page = fitz_doc[page_idx]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = pdf_page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # 添加分页符
            if i > 0:
                run = doc.add_paragraph().add_run()
                run.add_break(WD_BREAK.PAGE)

            # ---- 文字OCR识别 ----
            text_lines = None
            try:
                text_lines = client.recognize_text(img_bytes)
                logging.info(f"Page {page_num}: OCR recognized {len(text_lines)} lines")
            except Exception as e:
                err_msg = f"第{page_num}页OCR失败: {e}"
                logging.error(err_msg)
                ocr_errors.append(err_msg)

            if text_lines:
                for line_text in text_lines:
                    doc.add_paragraph(line_text)
            else:
                logging.info(f"Page {page_num}: No text recognized, inserting image")
                img_stream = io.BytesIO(img_bytes)
                page_width = pdf_page.rect.width / 72.0
                doc.add_picture(img_stream, width=Inches(min(page_width, 6.3)))

            # ---- 公式识别 ----
            if formula_api_on:
                try:
                    self._report(progress_text=f"识别第 {page_num} 页公式...")
                    formulas = client.recognize_formula(img_bytes)
                    for latex_str in formulas:
                        if not latex_str.strip():
                            continue
                        omml_elem = latex_to_omml(latex_str, xslt_path)
                        if omml_elem is not None:
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            insert_omml_to_paragraph(para, omml_elem)
                            formula_count += 1
                        else:
                            para = doc.add_paragraph(f"[公式] {latex_str}")
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            formula_count += 1
                except Exception as e:
                    logging.warning(f"Page {page_num} formula API error: {e}")

        doc.save(output_file)
        fitz_doc.close()

        result['formula_count'] = formula_count
        result['errors'] = ocr_errors
        self._report(percent=100, progress_text="转换完成！(100%)")

    @staticmethod
    def _ocr_mode_to_dpi(ocr_mode):
        mode = (ocr_mode or "平衡").strip()
        mapping = {
            "快速": 220,
            "平衡": 300,
            "高精": 360,
        }
        return mapping.get(mode, 300)

    # ----------------------------------------------------------
    # 公式后处理
    # ----------------------------------------------------------

    def _post_process_formula_api(self, docx_path, pdf_path, math_page_ids,
                                  client, xslt_path):
        """使用百度API识别公式并替换为Word原生OMML公式"""
        doc_obj = Document(docx_path)
        fitz_doc = fitz.open(pdf_path)
        fix_count = 0

        # 第一步：规范化数学Unicode字符
        for para in doc_obj.paragraphs:
            for run in para.runs:
                if has_math_unicode(run.text):
                    run.text = normalize_math_unicode(run.text)
                    fix_count += 1
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if has_math_unicode(run.text):
                                run.text = normalize_math_unicode(run.text)
                                fix_count += 1

        # 第二步：裁剪独立公式块 → API识别 → 替换为OMML
        for page_id in sorted(math_page_ids):
            pdf_page = fitz_doc[page_id]
            td = pdf_page.get_text("dict")
            for block in td.get("blocks", []):
                if not is_display_equation(block):
                    continue

                bbox = block["bbox"]
                x0, y0, x1, y1 = bbox
                padding = 5
                clip = fitz.Rect(
                    max(0, x0 - padding), max(0, y0 - padding),
                    min(pdf_page.rect.width, x1 + padding),
                    min(pdf_page.rect.height, y1 + padding)
                )
                if clip.is_empty or clip.width < 5 or clip.height < 5:
                    continue

                dpi = 300
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = pdf_page.get_pixmap(matrix=mat, clip=clip)
                img_bytes = pix.tobytes("png")

                try:
                    self._report(progress_text=f"正在识别第 {page_id + 1} 页的公式...")
                    formulas = client.recognize_formula(img_bytes)
                except Exception as e:
                    logging.warning(f"Formula API error on page {page_id + 1}: {e}")
                    continue

                if not formulas:
                    continue

                latex_str = formulas[0]
                block_text = get_block_text(block)
                norm_text = normalize_math_unicode(block_text)
                norm_compact = ''.join(norm_text.split())
                if len(norm_compact) < 2:
                    continue

                for para in doc_obj.paragraphs:
                    para_compact = ''.join(para.text.split())
                    if len(para_compact) < 2:
                        continue
                    if self._text_similar(para_compact, norm_compact):
                        omml_elem = latex_to_omml(latex_str, xslt_path)
                        if omml_elem is not None:
                            for run in para.runs:
                                run.text = ""
                            for child in list(para._element):
                                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                                if tag == 'r':
                                    para._element.remove(child)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            insert_omml_to_paragraph(para, omml_elem)
                            fix_count += 1
                            logging.info(f"Replaced equation with OMML: {latex_str[:50]}")
                        else:
                            for run in para.runs:
                                run.text = ""
                            img_stream = io.BytesIO(img_bytes)
                            region_width = clip.width / 72.0
                            doc_obj.add_picture(img_stream, width=Inches(min(region_width, 6.0)))
                            body = doc_obj.element.body
                            pic_element = body[-1]
                            para._element.addnext(pic_element)
                            fix_count += 1
                        break

        doc_obj.save(docx_path)
        fitz_doc.close()
        return fix_count

    # ----------------------------------------------------------
    # 文本相似度
    # ----------------------------------------------------------

    @staticmethod
    def _text_similar(a, b):
        """判断两个文本（已去空白）是否相似"""
        if not a or not b:
            return False
        if a == b:
            return True
        shorter = min(len(a), len(b))
        longer = max(len(a), len(b))
        if shorter < 3 or shorter / longer < 0.3:
            return False
        set_a, set_b = set(a), set(b)
        common_chars = set_a & set_b
        all_chars = set_a | set_b
        if not all_chars:
            return False
        jaccard = len(common_chars) / len(all_chars)
        if shorter >= 4 and (a[:shorter] in b or b[:shorter] in a):
            return True
        return jaccard > 0.6
